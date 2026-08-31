"""
document-conversion-api -- NEXUS candidate #6 (manual build, no FORGE).

Converts between PDF/Office documents and structured JSON, both directions,
using mature OSS libraries only (pdfplumber, python-docx, openpyxl, reportlab)
-- pure local computation, zero external network calls at runtime, no named
third-party data source. This sidesteps the BuyWhere-style hallucination risk
entirely (skills/asset-lifecycle): there is nothing external to hallucinate
about, the whole product is "run this well-known library over bytes the
caller sent us."

Every endpoint accepts base64-encoded file bytes directly in the JSON request
body -- never a URL to fetch. This is a deliberate scope boundary, not an
oversight: it removes SSRF from this asset's risk surface entirely (unlike
url-metadata-api/agent-verification-api, which both need SSRF guards because
they fetch caller-supplied URLs).

Built from manual_assets/agent-verification-api/main.py's structure (MCP +
x402 + FastAPI, one process, Cloud Run target) -- same telemetry helpers,
same x402 wiring, same discovery routes. Two things are genuinely new for
this asset (see the two docstring blocks below marked "NEW RISK"):

1. Every parse/generate call is synchronous and CPU-bound (pdfplumber/
   python-docx/openpyxl/reportlab have no async API), unlike the other 3
   manual/generated assets which are all I/O-bound via httpx.AsyncClient.
   Every handler offloads the actual work via asyncio.to_thread() inside an
   asyncio.wait_for() timeout -- inline synchronous calls in an async def
   route would block the whole event loop (every concurrent request) for the
   duration of one caller's parse.
2. .docx/.xlsx are zip archives -- a crafted file with a small compressed
   size but a huge uncompressed size (zip bomb) is a real DoS vector distinct
   from anything the other 3 assets face. _check_zip_bomb_safe() inspects the
   zip central directory (cheap, no decompression) before handing bytes to
   python-docx/openpyxl.
"""

import asyncio
import base64
import io
import os
import sys
import time
import zipfile
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from typing import Annotated, Literal, Optional

import httpx
from dotenv import load_dotenv

load_dotenv()

from fastapi import FastAPI, HTTPException, status
from fastapi.openapi.utils import get_openapi as _fastapi_get_openapi
from mcp.server.fastmcp import Context, FastMCP
from mcp.server.transport_security import TransportSecuritySettings
from pydantic import BaseModel, Field, field_validator, model_validator

from x402.http import FacilitatorConfig, HTTPFacilitatorClient, PaymentOption
from x402.http.middleware.fastapi import PaymentMiddlewareASGI
from x402.http.types import RouteConfig
from x402.mechanisms.evm.exact import ExactEvmServerScheme
from x402.schemas import Network
from x402.server import x402ResourceServer

_NEXUS_ASSET_NAME = "document-conversion-api"

# ---------------------------------------------------------------
# Size / resource caps -- the two "NEW RISK" categories from the module
# docstring. Every cap is checked BEFORE the expensive/unsafe operation it
# guards, never after.
# ---------------------------------------------------------------
_MAX_UPLOAD_BYTES = 8 * 1024 * 1024          # raw decoded upload size cap
_MAX_ZIP_UNCOMPRESSED_BYTES = 50 * 1024 * 1024  # docx/xlsx zip-bomb cap
_MAX_ZIP_COMPRESSION_RATIO = 100             # per-entry ratio cap (defense in depth)
_MAX_PDF_PAGES = 200                          # hard ceiling regardless of caller's max_pages
_PROCESSING_TIMEOUT_SECONDS = 25.0
_MAX_GENERATE_BLOCKS = 500
_MAX_TABLE_ROWS = 1000
_MAX_TABLE_COLS = 50
_MAX_BLOCK_TEXT_CHARS = 20000


def _decode_upload(file_base64: str) -> bytes:
    try:
        raw = base64.b64decode(file_base64, validate=True)
    except Exception:
        raise HTTPException(400, "file_base64 is not valid base64")
    if len(raw) > _MAX_UPLOAD_BYTES:
        raise HTTPException(
            413,
            f"Upload too large: {len(raw)} bytes decoded, cap is {_MAX_UPLOAD_BYTES} bytes.",
        )
    if not raw:
        raise HTTPException(400, "file_base64 decoded to zero bytes")
    return raw


def _check_zip_bomb_safe(raw: bytes) -> None:
    """
    NEW RISK (see module docstring, point 2). .docx/.xlsx are zip archives --
    reading the central directory (zipfile.infolist()) is cheap and does NOT
    decompress entry data, so this check runs before python-docx/openpyxl
    ever touch the bytes. Rejects a file whose total uncompressed size would
    exceed the cap, or whose any single entry has a compression ratio past
    _MAX_ZIP_COMPRESSION_RATIO (catches a small-total-but-pathological-entry
    bomb that a total-size check alone could miss).
    """
    try:
        zf = zipfile.ZipFile(io.BytesIO(raw))
        infos = zf.infolist()
    except zipfile.BadZipFile:
        raise HTTPException(400, "File is not a valid .docx/.xlsx (not a valid zip archive)")

    total_uncompressed = 0
    for info in infos:
        total_uncompressed += info.file_size
        if info.compress_size > 0:
            ratio = info.file_size / info.compress_size
            if ratio > _MAX_ZIP_COMPRESSION_RATIO:
                raise HTTPException(
                    400,
                    f"Rejected: entry {info.filename!r} has a suspicious compression "
                    f"ratio ({ratio:.0f}x) -- possible zip bomb.",
                )
    if total_uncompressed > _MAX_ZIP_UNCOMPRESSED_BYTES:
        raise HTTPException(
            400,
            f"Rejected: archive would decompress to {total_uncompressed} bytes, "
            f"cap is {_MAX_ZIP_UNCOMPRESSED_BYTES} bytes -- possible zip bomb.",
        )


_MAX_CONCURRENT_JOBS = int(os.getenv("NEXUS_MAX_CONCURRENT_JOBS", "4"))
_cpu_bound_semaphore = asyncio.Semaphore(_MAX_CONCURRENT_JOBS)


async def _run_cpu_bound(fn, *args):
    """
    NEW RISK (see module docstring, point 1). Every parser/generator call
    goes through here -- asyncio.to_thread() offloads the synchronous call
    to a worker thread so it never blocks the event loop, and
    asyncio.wait_for() bounds worst-case wall time regardless of what a
    pathological (but size-capped) input does inside the library.

    Known residual limitation: asyncio.wait_for() cancels the *awaiting*
    Task, it does not and cannot kill the underlying OS thread -- Python
    offers no API to forcibly terminate a running thread. A pathological
    input that hangs pdfminer/openpyxl internally keeps that one worker
    thread occupied indefinitely even after the caller gets a 504. The
    semaphore below bounds how many such "leaked" threads can accumulate
    concurrently (default 4) so a burst of adversarial calls degrades to
    503s for new requests rather than unbounded thread growth -- it does
    not reclaim an already-leaked thread. A full fix would need the CPU-
    bound work in a separate process (killable) rather than a thread;
    not done here as disproportionate for a 7-day probation candidate --
    see README "Known limitations".
    """
    try:
        await asyncio.wait_for(_cpu_bound_semaphore.acquire(), timeout=5.0)
    except asyncio.TimeoutError:
        raise HTTPException(503, "Too many concurrent conversions in progress, try again shortly")
    try:
        return await asyncio.wait_for(asyncio.to_thread(fn, *args), timeout=_PROCESSING_TIMEOUT_SECONDS)
    except asyncio.TimeoutError:
        raise HTTPException(504, f"Processing exceeded {_PROCESSING_TIMEOUT_SECONDS}s timeout")
    finally:
        _cpu_bound_semaphore.release()


# ---------------------------------------------------------------
# Supabase telemetry -- ported by hand from manual_assets/agent-verification-api
# (this asset is manual, not FORGE-generated). anon key only, never
# service_role -- RLS on traffic_events/mcp_call_events/revenue_events is
# INSERT-only for anon, the only real protection if this key leaks from
# the deployed runtime.
# ---------------------------------------------------------------
_NEXUS_SUPABASE_URL = os.getenv("SUPABASE_URL")
_NEXUS_SUPABASE_ANON_KEY = os.getenv("SUPABASE_ANON_KEY")

_nexus_bg_tasks: set = set()


def _nexus_fire_and_forget(coro):
    task = asyncio.create_task(coro)
    _nexus_bg_tasks.add(task)
    task.add_done_callback(_nexus_bg_tasks.discard)
    return task


def _nexus_truncate_ip(raw_ip: Optional[str]) -> Optional[str]:
    if not raw_ip:
        return None
    if ":" in raw_ip and "." not in raw_ip:
        segments = [s for s in raw_ip.split(":") if s]
        head = segments[:4] if len(segments) >= 4 else segments
        return (":".join(head) + "::/64") if head else None
    octets = raw_ip.split(".")
    if len(octets) == 4 and all(o.isdigit() for o in octets):
        return f"{octets[0]}.{octets[1]}.{octets[2]}.0/24"
    return None


async def _nexus_supabase_insert(table: str, payload: dict) -> None:
    if not _NEXUS_SUPABASE_URL or not _NEXUS_SUPABASE_ANON_KEY:
        return
    try:
        async with httpx.AsyncClient(timeout=3.0) as client:
            await client.post(
                f"{_NEXUS_SUPABASE_URL}/rest/v1/{table}",
                json=payload,
                headers={
                    "apikey": _NEXUS_SUPABASE_ANON_KEY,
                    "Authorization": f"Bearer {_NEXUS_SUPABASE_ANON_KEY}",
                    "Content-Type": "application/json",
                    "Prefer": "return=minimal",
                },
            )
    except Exception as e:
        print(f"[WARN] Supabase insert to {table!r} failed: {e}", file=sys.stderr)


async def _nexus_log_mcp_call_event(tool_id: str, success: bool, latency_ms: int) -> None:
    await _nexus_supabase_insert("mcp_call_events", {
        "tool_id": tool_id,
        "asset_name": _NEXUS_ASSET_NAME,
        "success": success,
        "latency_ms": latency_ms,
        "agent_framework": "mcp",
        "price_charged": 0,  # MCP path is free today -- see README "Known limitations"
    })


# =================================================================
# Extraction: document -> structured JSON
# =================================================================

class ExtractPdfRequest(BaseModel):
    file_base64: Annotated[str, Field(..., description="Base64-encoded PDF file bytes.")]
    max_pages: Annotated[int, Field(50, ge=1, le=_MAX_PDF_PAGES,
        description=f"Maximum pages to process. Hard ceiling {_MAX_PDF_PAGES} regardless of this value.")]


class ExtractDocxRequest(BaseModel):
    file_base64: Annotated[str, Field(..., description="Base64-encoded .docx file bytes.")]


class ExtractXlsxRequest(BaseModel):
    file_base64: Annotated[str, Field(..., description="Base64-encoded .xlsx file bytes.")]
    max_rows_per_sheet: Annotated[int, Field(1000, ge=1, le=_MAX_TABLE_ROWS)]
    max_cols_per_sheet: Annotated[int, Field(100, ge=1, le=_MAX_TABLE_COLS)]


def _extract_pdf_sync(raw: bytes, max_pages: int) -> dict:
    import pdfplumber

    pages_out = []
    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        total_pages = len(pdf.pages)
        cap = min(max_pages, _MAX_PDF_PAGES, total_pages)
        for i in range(cap):
            page = pdf.pages[i]
            text = page.extract_text() or ""
            tables = page.extract_tables() or []
            pages_out.append({
                "page_number": i + 1,
                "text": text,
                "tables": tables,
            })
        metadata = {k: v for k, v in (pdf.metadata or {}).items() if isinstance(v, str)}
    return {
        "page_count": total_pages,
        "pages_returned": len(pages_out),
        "pages": pages_out,
        "metadata": metadata,
        "truncated": total_pages > len(pages_out),
    }


def _extract_docx_sync(raw: bytes) -> dict:
    import docx

    d = docx.Document(io.BytesIO(raw))
    paragraphs = []
    for p in d.paragraphs:
        if not p.text.strip():
            continue
        style_name = p.style.name if p.style else None
        heading_level = None
        if style_name and style_name.lower().startswith("heading"):
            digits = "".join(c for c in style_name if c.isdigit())
            heading_level = int(digits) if digits else None
        paragraphs.append({"text": p.text, "style": style_name, "heading_level": heading_level})

    tables = []
    for t in d.tables:
        rows = [[cell.text for cell in row.cells] for row in t.rows]
        tables.append(rows)

    core = d.core_properties
    metadata = {
        "title": core.title,
        "author": core.author,
        "created": core.created.isoformat() if core.created else None,
    }
    return {"paragraphs": paragraphs, "tables": tables, "metadata": metadata}


def _extract_xlsx_sync(raw: bytes, max_rows: int, max_cols: int) -> dict:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    sheets = []
    for ws in wb.worksheets:
        rows = []
        truncated_rows = False
        truncated_cols = False
        for r_idx, row in enumerate(ws.iter_rows(values_only=True)):
            if r_idx >= max_rows:
                truncated_rows = True
                break
            row_vals = list(row)
            if len(row_vals) > max_cols:
                truncated_cols = True
                row_vals = row_vals[:max_cols]
            rows.append([str(v) if v is not None else None for v in row_vals])
        sheets.append({
            "name": ws.title,
            "rows": rows,
            "truncated": truncated_rows or truncated_cols,
        })
    wb.close()
    return {"sheets": sheets}


# =================================================================
# Generation: structured JSON -> document
# =================================================================

class GenerateBlock(BaseModel):
    type: Literal["heading", "paragraph", "table"]
    level: Annotated[Optional[int], Field(None, ge=1, le=6)] = None
    text: Annotated[Optional[str], Field(None, max_length=_MAX_BLOCK_TEXT_CHARS)] = None
    rows: Annotated[Optional[list[list[str]]], Field(None)] = None

    @field_validator("rows")
    @classmethod
    def _validate_rows(cls, v):
        if v is None:
            return v
        if len(v) > _MAX_TABLE_ROWS:
            raise ValueError(f"table has {len(v)} rows, cap is {_MAX_TABLE_ROWS}")
        for row in v:
            if len(row) > _MAX_TABLE_COLS:
                raise ValueError(f"table row has {len(row)} cols, cap is {_MAX_TABLE_COLS}")
        return v

    @model_validator(mode="after")
    def _validate_required_field_per_type(self):
        if self.type in ("heading", "paragraph") and not self.text:
            raise ValueError(f"block type {self.type!r} requires non-empty 'text'")
        if self.type == "table" and not self.rows:
            raise ValueError("block type 'table' requires non-empty 'rows'")
        return self


class GenerateDocRequest(BaseModel):
    title: Annotated[Optional[str], Field(None, max_length=500)] = None
    blocks: Annotated[list[GenerateBlock], Field(..., min_length=1, max_length=_MAX_GENERATE_BLOCKS)]


def _generate_pdf_sync(title: Optional[str], blocks: list) -> bytes:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER)
    styles = getSampleStyleSheet()
    story = []

    if title:
        story.append(Paragraph(title, styles["Title"]))
        story.append(Spacer(1, 0.25 * inch))

    for block in blocks:
        if block["type"] == "heading":
            level = block.get("level") or 1
            style_name = f"Heading{min(max(level, 1), 4)}"
            story.append(Paragraph(block.get("text") or "", styles.get(style_name, styles["Heading1"])))
        elif block["type"] == "paragraph":
            story.append(Paragraph(block.get("text") or "", styles["Normal"]))
        elif block["type"] == "table" and block.get("rows"):
            tbl = Table(block["rows"])
            tbl.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
            ]))
            story.append(tbl)
        story.append(Spacer(1, 0.15 * inch))

    doc.build(story)
    return buf.getvalue()


def _generate_docx_sync(title: Optional[str], blocks: list) -> bytes:
    import docx

    d = docx.Document()
    if title:
        d.add_heading(title, level=0)

    for block in blocks:
        if block["type"] == "heading":
            level = min(max(block.get("level") or 1, 1), 9)
            d.add_heading(block.get("text") or "", level=level)
        elif block["type"] == "paragraph":
            d.add_paragraph(block.get("text") or "")
        elif block["type"] == "table" and block.get("rows"):
            rows = block["rows"]
            n_cols = max((len(r) for r in rows), default=1)
            table = d.add_table(rows=0, cols=n_cols)
            for row in rows:
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = str(val) if val is not None else ""

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------
# MCP server (Streamable HTTP)
# ---------------------------------------------------------------
_public_domain = os.getenv("PUBLIC_DOMAIN", "*")
if _public_domain == "*":
    print(
        "[WARN] PUBLIC_DOMAIN is unset -- every real request will get 421'd once "
        "deployed publicly (fails closed). Set it to the real Cloud Run domain and "
        "redeploy.", file=sys.stderr,
    )

mcp = FastMCP(
    "document-conversion-api",
    stateless_http=True,
    host="0.0.0.0",
    streamable_http_path="/",
    transport_security=TransportSecuritySettings(
        enable_dns_rebinding_protection=True,
        allowed_hosts=["localhost:*", "127.0.0.1:*", _public_domain, _public_domain + ":*"],
        allowed_origins=["http://localhost:*", "http://127.0.0.1:*", "https://" + _public_domain],
    ),
)


@mcp.tool()
async def extract_pdf_to_json(file_base64: str, max_pages: int = 50, ctx: Context = None) -> dict:
    """Extract text, tables, page count and metadata from a PDF (base64-encoded bytes)."""
    start = time.monotonic()
    success = False
    try:
        raw = _decode_upload(file_base64)
        result = await _run_cpu_bound(_extract_pdf_sync, raw, max_pages)
        success = True
        return result
    finally:
        latency_ms = int((time.monotonic() - start) * 1000)
        _nexus_fire_and_forget(_nexus_log_mcp_call_event("extract_pdf_to_json", success, latency_ms))


@mcp.tool()
async def extract_docx_to_json(file_base64: str, ctx: Context = None) -> dict:
    """Extract paragraphs (with heading levels), tables and metadata from a .docx (base64-encoded bytes)."""
    start = time.monotonic()
    success = False
    try:
        raw = _decode_upload(file_base64)
        _check_zip_bomb_safe(raw)
        result = await _run_cpu_bound(_extract_docx_sync, raw)
        success = True
        return result
    finally:
        latency_ms = int((time.monotonic() - start) * 1000)
        _nexus_fire_and_forget(_nexus_log_mcp_call_event("extract_docx_to_json", success, latency_ms))


@mcp.tool()
async def extract_xlsx_to_json(file_base64: str, max_rows_per_sheet: int = 1000,
                                max_cols_per_sheet: int = 100, ctx: Context = None) -> dict:
    """Extract per-sheet cell grids from an .xlsx (base64-encoded bytes), capped rows/cols per sheet."""
    start = time.monotonic()
    success = False
    try:
        raw = _decode_upload(file_base64)
        _check_zip_bomb_safe(raw)
        result = await _run_cpu_bound(_extract_xlsx_sync, raw, max_rows_per_sheet, max_cols_per_sheet)
        success = True
        return result
    finally:
        latency_ms = int((time.monotonic() - start) * 1000)
        _nexus_fire_and_forget(_nexus_log_mcp_call_event("extract_xlsx_to_json", success, latency_ms))


@mcp.tool()
async def generate_pdf_from_json(title: Optional[str], blocks: list[dict], ctx: Context = None) -> dict:
    """Generate a PDF from structured blocks (heading/paragraph/table). Returns base64-encoded PDF bytes."""
    start = time.monotonic()
    success = False
    try:
        validated = GenerateDocRequest(title=title, blocks=blocks)
        pdf_bytes = await _run_cpu_bound(
            _generate_pdf_sync, validated.title, [b.model_dump() for b in validated.blocks]
        )
        success = True
        return {"file_base64": base64.b64encode(pdf_bytes).decode("ascii"), "byte_size": len(pdf_bytes)}
    finally:
        latency_ms = int((time.monotonic() - start) * 1000)
        _nexus_fire_and_forget(_nexus_log_mcp_call_event("generate_pdf_from_json", success, latency_ms))


@mcp.tool()
async def generate_docx_from_json(title: Optional[str], blocks: list[dict], ctx: Context = None) -> dict:
    """Generate a .docx from structured blocks (heading/paragraph/table). Returns base64-encoded docx bytes."""
    start = time.monotonic()
    success = False
    try:
        validated = GenerateDocRequest(title=title, blocks=blocks)
        docx_bytes = await _run_cpu_bound(
            _generate_docx_sync, validated.title, [b.model_dump() for b in validated.blocks]
        )
        success = True
        return {"file_base64": base64.b64encode(docx_bytes).decode("ascii"), "byte_size": len(docx_bytes)}
    finally:
        latency_ms = int((time.monotonic() - start) * 1000)
        _nexus_fire_and_forget(_nexus_log_mcp_call_event("generate_docx_from_json", success, latency_ms))


# ---------------------------------------------------------------
# FastAPI app
# ---------------------------------------------------------------
@asynccontextmanager
async def lifespan(app: FastAPI):
    async with mcp.session_manager.run():
        yield


app = FastAPI(
    title="Document Conversion API",
    description=(
        "Converts between PDF/Office documents and structured JSON, both directions. "
        "Pure OSS libraries (pdfplumber/python-docx/openpyxl/reportlab), zero external "
        "network calls at runtime -- every endpoint takes base64-encoded file bytes "
        "directly, never a URL."
    ),
    version="1.0.0",
    contact={"email": "dasaanrod@gmail.com"},
    lifespan=lifespan,
)


@app.middleware("http")
async def _nexus_traffic_log(request, call_next):
    response = await call_next(request)
    ip = request.client.host if request.client else None
    _nexus_fire_and_forget(_nexus_supabase_insert("traffic_events", {
        "asset_name": _NEXUS_ASSET_NAME,
        "ip_range": _nexus_truncate_ip(ip),
        "method": request.method,
        "path": request.url.path,
        "status": response.status_code,
    }))
    return response


# --- x402: pay-per-call in USDC, Base Sepolia testnet -- same wallet,
#     facilitator and self-payment-bug fix as the sibling manual assets
#     (skills/x402-payments). Pricing: no paid third-party API cost here
#     (unlike WHOIS in live-entity-verification/agent-verification-api) --
#     pure CPU/memory cost, priced in the same low tier as url-metadata-api's
#     $0.01 single-fetch. PDF operations (extract/generate) cost a notch more
#     -- pdfplumber/reportlab do meaningfully more work per call than
#     python-docx/openpyxl (page rendering vs. straight XML parsing). ---
_NEXUS_X402_FREE_MODE = os.getenv("NEXUS_X402_FREE_MODE", "false").strip().lower() == "true"

_X402_EVM_ADDRESS = os.getenv("X402_WALLET_ADDRESS", "0xYOUR_WALLET_ADDRESS_HERE")
_X402_NETWORK: Network = "eip155:84532"  # Base Sepolia testnet

_looks_like_evm_address = (
    _X402_EVM_ADDRESS.startswith("0x")
    and len(_X402_EVM_ADDRESS) == 42
    and all(c in "0123456789abcdefABCDEF" for c in _X402_EVM_ADDRESS[2:])
)
if not _NEXUS_X402_FREE_MODE and not _looks_like_evm_address:
    print(
        f"[WARN] X402_WALLET_ADDRESS ({_X402_EVM_ADDRESS!r}) doesn't look like a "
        "real EVM address (expected '0x' + 40 hex chars). The server will still "
        "boot and issue 402 challenges, but payments will have nowhere real to "
        "settle. See README.md.",
        file=sys.stderr,
    )

_x402_facilitator = HTTPFacilitatorClient(FacilitatorConfig(url="https://x402.org/facilitator"))
_x402_server = x402ResourceServer(_x402_facilitator)
_x402_server.register(_X402_NETWORK, ExactEvmServerScheme())


async def _nexus_log_x402_revenue_event(ctx) -> None:
    try:
        result = getattr(ctx, "result", None)
        requirements = getattr(ctx, "requirements", None)
        if result is None or requirements is None or not getattr(result, "success", False):
            return
        raw_amount = getattr(requirements, "amount", None)
        amount_eur = int(raw_amount) / 1_000_000 if raw_amount is not None else None
        await _nexus_supabase_insert("revenue_events", {
            "asset_name": _NEXUS_ASSET_NAME,
            "amount_eur": amount_eur,
            "pricing_model": "x402",
            "stripe_event_id": None,
            "customer_id": getattr(result, "payer", None),
        })
    except Exception:
        pass


_x402_server.on_after_settle(_nexus_log_x402_revenue_event)

_X402_ROUTE_PRICES = {
    "POST /extract-pdf-to-json": "$0.02",
    "POST /extract-docx-to-json": "$0.01",
    "POST /extract-xlsx-to-json": "$0.01",
    "POST /generate-pdf-from-json": "$0.02",
    "POST /generate-docx-from-json": "$0.01",
}

_X402_ROUTES: dict[str, RouteConfig] = {
    route_key: RouteConfig(
        accepts=[PaymentOption(scheme="exact", pay_to=_X402_EVM_ADDRESS, price=price, network=_X402_NETWORK)],
        mime_type="application/json",
        description=route_key,
    )
    for route_key, price in _X402_ROUTE_PRICES.items()
}
if not _NEXUS_X402_FREE_MODE:
    app.add_middleware(PaymentMiddlewareASGI, routes=_X402_ROUTES, server=_x402_server)


@app.post(
    "/extract-pdf-to-json",
    responses={
        400: {"description": "invalid base64, empty upload, or oversized upload"},
        504: {"description": f"processing exceeded {_PROCESSING_TIMEOUT_SECONDS}s"},
    },
)
async def extract_pdf_to_json_endpoint(payload: ExtractPdfRequest) -> dict:
    """Extract text, tables (per page), page count and metadata from a PDF.
    A 400 (invalid base64/oversized upload) or 504 (processing timeout) raised here is NOT
    charged -- corrected 2026-08-31, see main NEXUS repo CLAUDE.md SS8: x402's
    PaymentMiddlewareASGI only settles after a <400 response, verified empirically."""
    raw = _decode_upload(payload.file_base64)
    return await _run_cpu_bound(_extract_pdf_sync, raw, payload.max_pages)


@app.post(
    "/extract-docx-to-json",
    responses={
        400: {"description": "invalid base64, empty/oversized upload, not a valid zip, or zip-bomb rejected"},
        504: {"description": f"processing exceeded {_PROCESSING_TIMEOUT_SECONDS}s"},
    },
)
async def extract_docx_to_json_endpoint(payload: ExtractDocxRequest) -> dict:
    """Extract paragraphs (with heading levels/styles) and tables from a .docx.
    Same not-charged-on-failure correction as /extract-pdf-to-json (2026-08-31)."""
    raw = _decode_upload(payload.file_base64)
    _check_zip_bomb_safe(raw)
    return await _run_cpu_bound(_extract_docx_sync, raw)


@app.post(
    "/extract-xlsx-to-json",
    responses={
        400: {"description": "invalid base64, empty/oversized upload, not a valid zip, or zip-bomb rejected"},
        504: {"description": f"processing exceeded {_PROCESSING_TIMEOUT_SECONDS}s"},
    },
)
async def extract_xlsx_to_json_endpoint(payload: ExtractXlsxRequest) -> dict:
    """Extract per-sheet cell grids from an .xlsx, capped rows/cols per sheet.
    Same not-charged-on-failure correction as /extract-pdf-to-json (2026-08-31)."""
    raw = _decode_upload(payload.file_base64)
    _check_zip_bomb_safe(raw)
    return await _run_cpu_bound(_extract_xlsx_sync, raw, payload.max_rows_per_sheet, payload.max_cols_per_sheet)


@app.post(
    "/generate-pdf-from-json",
    responses={
        422: {"description": "invalid block structure (bad type, oversized table/text)"},
        504: {"description": f"processing exceeded {_PROCESSING_TIMEOUT_SECONDS}s"},
    },
)
async def generate_pdf_from_json_endpoint(payload: GenerateDocRequest) -> dict:
    """Generate a PDF from structured blocks (heading/paragraph/table). Returns
    base64-encoded PDF bytes. Same not-charged-on-failure correction as
    /extract-pdf-to-json (2026-08-31) -- a 422 (malformed blocks) is NOT charged either."""
    pdf_bytes = await _run_cpu_bound(_generate_pdf_sync, payload.title, [b.model_dump() for b in payload.blocks])
    return {"file_base64": base64.b64encode(pdf_bytes).decode("ascii"), "byte_size": len(pdf_bytes)}


@app.post(
    "/generate-docx-from-json",
    responses={
        422: {"description": "invalid block structure (bad type, oversized table/text)"},
        504: {"description": f"processing exceeded {_PROCESSING_TIMEOUT_SECONDS}s"},
    },
)
async def generate_docx_from_json_endpoint(payload: GenerateDocRequest) -> dict:
    """Generate a .docx from structured blocks (heading/paragraph/table). Returns
    base64-encoded docx bytes. Same not-charged-on-failure correction as
    /extract-pdf-to-json (2026-08-31)."""
    docx_bytes = await _run_cpu_bound(_generate_docx_sync, payload.title, [b.model_dump() for b in payload.blocks])
    return {"file_base64": base64.b64encode(docx_bytes).decode("ascii"), "byte_size": len(docx_bytes)}


@app.get("/health")
async def health() -> dict:
    return {"status": "ok"}


_FAVICON_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


@app.get("/favicon.ico", include_in_schema=False)
async def favicon():
    from fastapi.responses import Response
    return Response(content=_FAVICON_PNG, media_type="image/png")


# --- 402index.io domain claim verification (POST /api/v1/claim -- filled in
#     after the real Cloud Run domain is known, see README "Deploy") ---
_NEXUS_402INDEX_VERIFY_HASH = os.getenv("NEXUS_402INDEX_VERIFY_HASH", "")


@app.get("/.well-known/402index-verify.txt", include_in_schema=False)
async def _nexus_402index_verify():
    from fastapi.responses import PlainTextResponse
    if not _NEXUS_402INDEX_VERIFY_HASH:
        raise HTTPException(404, "no 402index claim on file")
    return PlainTextResponse(content=_NEXUS_402INDEX_VERIFY_HASH)


@app.get("/.well-known/agent-card.json", include_in_schema=False)
async def agent_card() -> dict:
    base = f"https://{_public_domain}" if _public_domain != "*" else "https://document-conversion-api.example"
    return {
        "name": "Document Conversion API",
        "description": "Converts between PDF/Office documents and structured JSON, both directions. "
                        "Pure OSS libraries, zero external network calls at runtime.",
        "url": base,
        "version": "1.0.0",
        "documentationUrl": f"{base}/docs",
        "provider": {"organization": "nexus-mcp-infra", "url": "https://github.com/nexus-mcp-infra"},
        "capabilities": {"streaming": False, "pushNotifications": False, "stateTransitionHistory": False},
        "defaultInputModes": ["application/json"],
        "defaultOutputModes": ["application/json"],
        "additionalInterfaces": [{"url": f"{base}/mcp", "transport": "MCP"}],
        "skills": [
            {
                "id": "nexus_document_conversion_api_extract_pdf_to_json",
                "name": "Extract PDF to JSON",
                "description": "Extract text, tables, page count and metadata from a PDF.",
                "tags": ["pdf", "extraction", "document-conversion"],
            },
            {
                "id": "nexus_document_conversion_api_extract_docx_to_json",
                "name": "Extract DOCX to JSON",
                "description": "Extract paragraphs, headings and tables from a .docx.",
                "tags": ["docx", "office", "extraction"],
            },
            {
                "id": "nexus_document_conversion_api_extract_xlsx_to_json",
                "name": "Extract XLSX to JSON",
                "description": "Extract per-sheet cell grids from an .xlsx.",
                "tags": ["xlsx", "office", "extraction"],
            },
            {
                "id": "nexus_document_conversion_api_generate_pdf_from_json",
                "name": "Generate PDF from JSON",
                "description": "Generate a PDF from structured blocks (heading/paragraph/table).",
                "tags": ["pdf", "generation", "document-conversion"],
            },
            {
                "id": "nexus_document_conversion_api_generate_docx_from_json",
                "name": "Generate DOCX from JSON",
                "description": "Generate a .docx from structured blocks (heading/paragraph/table).",
                "tags": ["docx", "office", "generation"],
            },
        ],
        "metadata": {
            "protocol_note": (
                "This service implements the Model Context Protocol (MCP) at /mcp, not A2A's "
                "own task methods. Every business POST route is charged via x402 (Base Sepolia "
                "TESTNET, not real funds): $0.02 for PDF operations (extract/generate), $0.01 "
                "for docx/xlsx operations. Payment only settles after a successful (<400) "
                "response: a call that returns 400 (invalid/oversized/zip-bomb-rejected upload), "
                "422 (malformed generate blocks), or 504 (processing timeout) is NOT charged. The "
                "MCP tool surface at /mcp is currently free (in-process call, not re-metered)."
            ),
        },
    }


_NEXUS_X402_OPENAPI_OPERATIONS = [(m.split(" ")[0].lower(), m.split(" ")[1]) for m in _X402_ROUTE_PRICES]


def _nexus_openapi_with_payment_info():
    if app.openapi_schema:
        return app.openapi_schema
    schema = _fastapi_get_openapi(title=app.title, version=app.version, description=app.description, routes=app.routes)
    if not _NEXUS_X402_FREE_MODE:
        for method, path in _NEXUS_X402_OPENAPI_OPERATIONS:
            op = schema.get("paths", {}).get(path, {}).get(method)
            if op is None:
                continue
            price = _X402_ROUTE_PRICES[f"{method.upper()} {path}"]
            op["x-payment-info"] = {
                "price": {"mode": "fixed", "currency": "USD", "amount": price.lstrip("$")},
                "protocols": [{"x402": {}}],
            }
            op.setdefault("responses", {})["402"] = {"description": "Payment Required"}
    schema.setdefault("info", {})["contact"] = {"email": "dasaanrod@gmail.com"}
    app.openapi_schema = schema
    return app.openapi_schema


app.openapi = _nexus_openapi_with_payment_info

app.mount("/mcp", mcp.streamable_http_app())
